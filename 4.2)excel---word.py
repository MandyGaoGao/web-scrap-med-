# -*- coding: utf-8 -*-
"""
Created on Thu Jul  5 16:25:40 2018

@author: gaoyu
"""

from docx import Document
import pandas as pd
import time
#get dic from word (done)
#get word from dic,formatting in a function
start_time = time.time()

def getword(name,ls):
    document = Document()
    document.add_heading(name, 0)
    p = document.add_paragraph('Composition en substances actives:')
    p.add_run(ls[1]).bold = True
    p.add_run(' Présentations:')
    p.add_run(ls[0]).italic = True

    document.add_heading('RCP', level=1)
    document.add_paragraph(ls[2], style='IntenseQuote')

    document.add_paragraph(ls[3])

    document.add_page_break()

    document.save('{}.docx'.format(name.replace('/',' ')))
    
    
def Main():
    a=pd.read_excel('RCP testing.xlsx')
    namelst=[]
    ls=[]
    names=a.columns
    for name in names:
        namelst.append(name)
    
    for name in namelst:
       
        ls=[]
        for i in a.index:
            ls.append(a[product][i])
        
        getword(name,ls)
        
if __name__=='__main__':
    Main()
print("--- %s seconds ---" % (time.time() - start_time))      